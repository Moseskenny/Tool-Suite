import os
import subprocess
import shutil
import io
import markdown
import fitz
from pdf2docx import Converter
import img2pdf
from docx import Document
from PIL import Image


def _prepare_output_dir(output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    for f in os.listdir(output_dir):
        os.remove(os.path.join(output_dir, f))


def _single(handler):
    def wrapped(input_path, output_dir, base_name, target_format):
        output_path = os.path.join(output_dir, f"converted_{base_name}.{target_format}")
        handler(input_path, output_path)
        return [f"converted_{base_name}.{target_format}"]
    return wrapped


# ---- Individual conversion handlers ----

def pdf_to_docx(input_path, output_path):
    cv = Converter(input_path)
    cv.convert(output_path)
    cv.close()


def pdf_to_txt(input_path, output_path):
    doc = fitz.open(input_path)
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


def pdf_to_images_handler(input_path, output_dir, base_name, target_format):
    doc = fitz.open(input_path)
    filenames = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        fname = f"converted_{base_name}_page_{i+1}.{target_format}"
        pix.save(os.path.join(output_dir, fname))
        filenames.append(fname)
    doc.close()
    return filenames


def docx_to_pdf_handler(input_path, output_dir, base_name, target_format):
    output_path = os.path.join(output_dir, f"converted_{base_name}.{target_format}")

    soffice_path = shutil.which("soffice")
    if soffice_path:
        try:
            subprocess.run(
                [soffice_path, "--headless", "--convert-to", "pdf",
                 "--outdir", output_dir, os.path.abspath(input_path)],
                check=True, capture_output=True, timeout=30
            )
            lo_name = base_name + ".pdf"
            lo_path = os.path.join(output_dir, lo_name)
            if os.path.exists(lo_path) and lo_path != output_path:
                shutil.move(lo_path, output_path)
                return [f"converted_{base_name}.{target_format}"]
        except Exception:
            pass

    _docx_to_pdf_fallback(input_path, output_path)
    return [f"converted_{base_name}.{target_format}"]


def _docx_to_pdf_fallback(input_path, output_path):
    from fpdf import FPDF
    doc = Document(input_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for para in doc.paragraphs:
        text = para.text[:2000].encode("latin-1", "replace").decode("latin-1")
        if text.strip():
            pdf.multi_cell(0, 8, text)
            pdf.ln(2)
    pdf.output(output_path)


def docx_to_txt(input_path, output_path):
    doc = Document(input_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)


def md_to_docx(input_path, output_path):
    doc = Document()
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")
    for line in lines:
        s = line.strip()
        if s.startswith("### "):
            doc.add_heading(s[4:], 3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], 1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s:
            doc.add_paragraph(s)
    doc.save(output_path)


def html_to_pdf(input_path, output_path):
    from weasyprint import HTML
    with open(input_path, "r", encoding="utf-8") as f:
        html_str = f.read()
    HTML(string=html_str).write_pdf(output_path)


def txt_to_pdf(input_path, output_path):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    with open(input_path, "r", encoding="utf-8") as f:
        for line in f:
            safe = line.encode("latin-1", "replace").decode("latin-1")
            pdf.cell(200, 10, txt=safe, ln=1)
    pdf.output(output_path)


def txt_to_docx(input_path, output_path):
    doc = Document()
    with open(input_path, "r", encoding="utf-8") as f:
        doc.add_paragraph(f.read())
    doc.save(output_path)


def md_to_pdf(input_path, output_path):
    from fpdf import FPDF
    with open(input_path, "r", encoding="utf-8") as f:
        html = markdown.markdown(f.read())
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 10, txt=html.encode("latin-1", "replace").decode("latin-1"))
    pdf.output(output_path)


def image_to_pdf(input_path, output_path):
    with open(output_path, "wb") as f:
        f.write(img2pdf.convert(input_path))


def image_convert_handler(input_path, output_dir, base_name, target_format):
    output_path = os.path.join(output_dir, f"converted_{base_name}.{target_format}")
    img = Image.open(input_path).convert("RGB")
    img.save(output_path, target_format.upper())
    return [f"converted_{base_name}.{target_format}"]


# ---- Dispatch table ----

HANDLERS = {
    ('.pdf', 'docx'): _single(pdf_to_docx),
    ('.pdf', 'txt'): _single(pdf_to_txt),
    ('.pdf', 'jpg'): pdf_to_images_handler,
    ('.pdf', 'png'): pdf_to_images_handler,
    ('.docx', 'pdf'): docx_to_pdf_handler,
    ('.docx', 'txt'): _single(docx_to_txt),
    ('.md', 'pdf'): _single(md_to_pdf),
    ('.md', 'docx'): _single(md_to_docx),
    ('.html', 'pdf'): _single(html_to_pdf),
    ('.txt', 'pdf'): _single(txt_to_pdf),
    ('.txt', 'docx'): _single(txt_to_docx),
    ('.jpg', 'pdf'): _single(image_to_pdf),
    ('.jpeg', 'pdf'): _single(image_to_pdf),
    ('.png', 'pdf'): _single(image_to_pdf),
    ('.jpg', 'png'): image_convert_handler,
    ('.jpg', 'jpeg'): image_convert_handler,
    ('.jpeg', 'png'): image_convert_handler,
    ('.png', 'jpg'): image_convert_handler,
    ('.png', 'jpeg'): image_convert_handler,
    ('.jpg', 'webp'): image_convert_handler,
    ('.jpeg', 'webp'): image_convert_handler,
    ('.png', 'webp'): image_convert_handler,
}


def run_conversion(files, target_format, output_dir, progress_callback=None):
    try:
        _prepare_output_dir(output_dir)

        total = len(files)
        all_results = []
        for idx, file in enumerate(files):
            if progress_callback:
                progress_callback(int((idx) / total * 100))

            input_path = os.path.join(output_dir, file.filename)
            with open(input_path, "wb") as f:
                f.write(file.file.read())

            ext = os.path.splitext(file.filename)[1].lower()
            base = os.path.splitext(file.filename)[0]
            key = (ext, target_format)

            if key not in HANDLERS:
                if os.path.exists(input_path):
                    os.remove(input_path)
                return {"status": "error", "message": f"Unsupported conversion: {ext} \u2192 .{target_format}"}

            try:
                handler = HANDLERS[key]
                filenames = handler(input_path, output_dir, base, target_format)
                all_results.extend(filenames)
            except Exception as e:
                if os.path.exists(input_path):
                    os.remove(input_path)
                return {"status": "error", "message": f"Failed to convert {file.filename}: {str(e)}"}

            if os.path.exists(input_path):
                os.remove(input_path)

            if progress_callback:
                progress_callback(int((idx + 1) / total * 100))

        remaining = [f for f in os.listdir(output_dir) if os.path.isfile(os.path.join(output_dir, f))]
        sizes = {f: os.path.getsize(os.path.join(output_dir, f)) for f in remaining}
        print(f"[Converter] Output dir: {output_dir}")
        print(f"[Converter] Files: {sizes}")

        return {
            "status": "success",
            "message": f"Successfully converted {len(files)} file(s).",
            "files": all_results,
            "file_count": len(all_results),
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


def analyze_files(files, output_dir):
    results = []
    for file in files:
        content = file.file.read()
        ext = os.path.splitext(file.filename)[1].lower()
        info = {
            "name": file.filename,
            "size": len(content),
            "ext": ext,
        }
        try:
            if ext == '.pdf':
                doc = fitz.open(stream=content, filetype="pdf")
                info["pages"] = len(doc)
                doc.close()
            elif ext == '.docx':
                doc = Document(io.BytesIO(content))
                info["paragraphs"] = len(doc.paragraphs)
            elif ext in ['.jpg', '.jpeg', '.png']:
                img = Image.open(io.BytesIO(content))
                info["width"] = img.width
                info["height"] = img.height
            elif ext in ['.txt', '.md']:
                text = content.decode("utf-8", errors="replace")
                info["lines"] = len(text.splitlines())
                info["chars"] = len(text)
        except Exception:
            pass
        results.append(info)
    return {"status": "success", "files": results}
